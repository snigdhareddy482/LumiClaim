"""Generate appeal documents grounded in existing explainability artifacts."""

from __future__ import annotations

import hashlib
import json
from datetime import datetime, timezone
from io import BytesIO
from typing import Any

try:
    from docx import Document  # type: ignore[import-not-found]
    _HAS_DOCX = True
except Exception:  # pragma: no cover - optional dependency
    Document: Any = None
    _HAS_DOCX = False

try:
    from fpdf import FPDF, XPos, YPos  # type: ignore[import-not-found]
    from fpdf.errors import FPDFException
    _HAS_FPDF = True
except Exception:  # pragma: no cover - optional dependency
    # Minimal fallback implementations so static type checkers and
    # import-time checks succeed even when fpdf2 isn't installed.
    class _FPDFFallback:
        def __init__(self, *args, **kwargs) -> None:
            self.font_family = "Helvetica"
            self.font_style = ""
            self.font_size_pt = 12

        def set_auto_page_break(self, auto: bool, margin: int) -> None:
            return None

        def add_page(self) -> None:
            return None

        def set_font(self, family: str, style: str = "", size: int = 0) -> None:
            self.font_family = family
            self.font_style = style
            self.font_size_pt = size

        def cell(self, *args, **kwargs) -> None:
            return None

        def ln(self, h: int = 0) -> None:
            return None

        def multi_cell(self, w: int, h: int, txt: str) -> None:
            return None

        def set_y(self, y: int) -> None:
            return None

        def page_no(self) -> int:
            return 1

        def output(self, dest: str = "S") -> bytes:
            return b""

    class _XPos:
        LMARGIN = 0

    class _YPos:
        NEXT = 0

    FPDF = _FPDFFallback  # type: ignore[assignment]
    XPos = _XPos  # type: ignore[assignment]
    YPos = _YPos  # type: ignore[assignment]
    FPDFException = Exception  # type: ignore[assignment]
    _HAS_FPDF = False

from backend.egraph import build_evidence_graph
from backend.math_guard import explain_bill
from backend.session import record_audit_entry

_TONE_TEMPLATES = {
    "polite": {
        "opening": "I am writing to request a careful review of the enclosed claim determination.",
        "closing": "Thank you for your time and prompt attention to this matter.",
    },
    "firm": {
        "opening": (
            "This letter lodges a formal dispute of the enclosed claim outcome based on the attached math evidence."
        ),
        "closing": "Please respond with a corrective action plan within the timeframe required by regulation.",
    },
}

_AUDIENCE_GREETINGS = {
    "payer": "To the Appeals Team",
    "provider": "To the Provider Billing Office",
}


def _build_context(
    doc_id: str,
    tone: str,
    audience: str,
    session_id: str | None,
) -> dict[str, Any]:
    explain_payload = explain_bill(doc_id, session_id=session_id)
    evidence_graph = build_evidence_graph(doc_id, session_id=session_id)
    selected_tone = _TONE_TEMPLATES.get(tone, _TONE_TEMPLATES["polite"])
    greeting = _AUDIENCE_GREETINGS.get(audience, _AUDIENCE_GREETINGS["payer"])
    breakdown_summary = _summarize_breakdown(explain_payload.get("breakdown", []) or [])
    risk_summary = _summarize_risks(explain_payload.get("risk_flags", []) or [])

    return {
        "doc_id": doc_id,
        "subject": f"Appeal for {doc_id}",
        "tone": tone,
        "audience": audience,
        "selected_tone": selected_tone,
        "greeting": greeting,
        "breakdown_summary": breakdown_summary,
        "risk_summary": risk_summary,
        "explain_payload": explain_payload,
        "evidence_graph": evidence_graph,
    }


def _fmt_currency(value: float | None) -> str:
    if value is None:
        return "an unknown amount"
    abs_value = abs(value)
    formatted = f"${abs_value:,.2f}"
    return formatted if value >= 0 else f"-{formatted}"


def _summarize_breakdown(breakdown: list[dict[str, Any]]) -> list[str]:
    lines: list[str] = []
    for item in breakdown:
        label = item.get("label", "Line")
        value_raw = item.get("value")
        value = float(value_raw) if value_raw is not None else None
        lines.append(f"{label}: {_fmt_currency(value)}")
    return lines


def _summarize_risks(risk_flags: list[dict[str, Any]]) -> list[str]:
    lines: list[str] = []
    for flag in risk_flags:
        label = flag.get("label", "Risk")
        severity = flag.get("severity", "unknown").title()
        rationale = flag.get("rationale", "No rationale provided.")
        lines.append(f"{label} ({severity}): {rationale}")
    return lines


def _build_exhibit_index(has_risk: bool) -> list[dict[str, str]]:
    exhibit_index = [
        {
            "label": "A",
            "title": "PME Explainability Packet",
            "description": "Breakdown, calculations, warnings, and citations used in this appeal.",
        },
        {
            "label": "B",
            "title": "Evidence Graph Snapshot",
            "description": "Graph of amounts, codes, sources, and relations supporting the claim.",
        },
    ]

    if has_risk:
        exhibit_index.append(
            {
                "label": "C",
                "title": "Risk Assessment Notes",
                "description": "Heuristic risk flags drawn from claim line analysis.",
            }
        )

    return exhibit_index


def build_appeal(
    doc_id: str,
    *,
    tone: str = "polite",
    audience: str = "payer",
    psl_delta: float | None = None,
    session_id: str | None = None,
) -> dict[str, Any]:
    """Construct a structured appeals packet referencing prior analysis."""

    context = _build_context(doc_id, tone, audience, session_id)
    explain_payload = context["explain_payload"]
    evidence_graph = context["evidence_graph"]
    selected_tone = context["selected_tone"]
    greeting = context["greeting"]
    breakdown_summary = context["breakdown_summary"]
    risk_summary = context["risk_summary"]

    body_sections: list[str] = [
        f"{greeting},",
        "",
        selected_tone["opening"],
    ]
    body_sections.append("")
    body_sections.append("Primary math evidence (PME) summary:")

    if breakdown_summary:
        for line in breakdown_summary:
            body_sections.append(f"- {line}")
    else:
        body_sections.append("- No structured totals available.")

    if risk_summary:
        body_sections.append("")
        body_sections.append("Risk considerations:")
        for line in risk_summary:
            body_sections.append(f"- {line}")

    if psl_delta is not None:
        body_sections.append("")
        body_sections.append("Policy simulation variance:")
        body_sections.append(
            f"- Expected vs billed patient responsibility delta: {_fmt_currency(psl_delta)}"
        )

    body_sections.extend([
        "",
        selected_tone["closing"],
        "",
        "Sincerely,",
        "Patient Advocate",
    ])

    body = "\n".join(body_sections)

    exhibit_index = _build_exhibit_index(bool(risk_summary))

    proof_pack = {
        "doc_id": doc_id,
        "explain": explain_payload,
        "evidence_graph": evidence_graph,
        "exhibit_index": exhibit_index,
    }

    exhibits = [
        {
            "label": item["label"],
            "title": item["title"],
            "refs": "proof_pack",
        }
        for item in exhibit_index
    ]

    payload = {
        "doc_id": doc_id,
        "tone": tone,
        "audience": audience,
        "subject": context["subject"],
        "body": body,
        "exhibits": exhibits,
        "psl_delta": psl_delta,
        "proof_pack": proof_pack,
    }

    audit_hash = hashlib.sha256(json.dumps(payload, sort_keys=True).encode("utf-8")).hexdigest()[:16]
    payload["audit_hash"] = audit_hash

    if session_id:
        try:
            record_audit_entry(session_id, "appeal", payload)
        except Exception:
            # Auditing is best-effort; ignore persistence issues
            pass

    return payload


def build_appeal_docx(
    doc_id: str,
    *,
    tone: str = "polite",
    audience: str = "payer",
    psl_delta: float | None = None,
    session_id: str | None = None,
) -> bytes:
    """Generate a DOCX representation of the appeal letter."""

    if not _HAS_DOCX:  # pragma: no cover - requires python-docx
        raise RuntimeError(
            "python-docx is required to build a DOCX appeal. Install python-docx and try again."
        )

    context = _build_context(doc_id, tone, audience, session_id)
    breakdown_summary = context["breakdown_summary"]
    risk_summary = context["risk_summary"]
    selected_tone = context["selected_tone"]
    greeting = context["greeting"]

    document = Document()
    document.add_heading(context["subject"], level=1)

    document.add_paragraph(f"Tone: {tone.title()} | Audience: {audience.title()}")
    document.add_paragraph("")
    document.add_paragraph(f"{greeting},")
    document.add_paragraph(selected_tone["opening"])

    document.add_heading("Primary math evidence (PME) summary", level=2)
    if breakdown_summary:
        for line in breakdown_summary:
            document.add_paragraph(line, style="List Bullet")
    else:
        document.add_paragraph("No structured totals available.")

    if risk_summary:
        document.add_heading("Risk considerations", level=2)
        for line in risk_summary:
            document.add_paragraph(line, style="List Bullet")

    if psl_delta is not None:
        document.add_heading("Policy simulation variance", level=2)
        document.add_paragraph(
            f"Expected vs billed patient responsibility delta: {_fmt_currency(psl_delta)}"
        )

    document.add_paragraph(selected_tone["closing"])
    document.add_paragraph("")
    document.add_paragraph("Sincerely,")
    document.add_paragraph("Patient Advocate")

    exhibit_index = _build_exhibit_index(bool(risk_summary))
    document.add_heading("Exhibits", level=2)
    for exhibit in exhibit_index:
        label = exhibit["label"]
        title = exhibit["title"]
        description = exhibit.get("description", "")
        document.add_paragraph(f"{label}. {title} - {description}", style="List Number")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


# Ensure FPDF is a class at runtime so static analyzers and a Python
# compile-time pass will accept using it as a base class.
if not isinstance(FPDF, type):  # pragma: no cover - fallback for static analysis
    class _FPDFCompat:
        def __init__(self, *args, **kwargs) -> None:
            self.font_family = "Helvetica"
            self.font_style = ""
            self.font_size_pt = 12

        def set_auto_page_break(self, auto: bool, margin: int) -> None:
            return None

        def add_page(self) -> None:
            return None

        def set_font(self, family: str, style: str = "", size: int = 0) -> None:
            self.font_family = family
            self.font_style = style
            self.font_size_pt = size

        def cell(self, *args, **kwargs) -> None:
            return None

        def ln(self, h: int = 0) -> None:
            return None

        def multi_cell(self, w: int, h: int, txt: str) -> None:
            return None

        def set_y(self, y: int) -> None:
            return None

        def page_no(self) -> int:
            return 1

        def output(self, dest: str = "S") -> bytes:
            return b""

    FPDF = _FPDFCompat  # type: ignore[assignment]

# Select a concrete base class before defining the class to avoid using an
# inline conditional expression directly in the class header.
_BasePDF = FPDF if isinstance(FPDF, type) else object

class _AppealPDF(_BasePDF):  # type: ignore[misc]
    def __init__(self, *, doc_id: str, generated_at: str) -> None:
        super().__init__()
        self.doc_id = doc_id
        self.generated_at = generated_at

    def header(self) -> None:  # type: ignore[override]
        self.set_font("Helvetica", "B", 12)
        # use new positioning args (new_x/new_y) rather than deprecated ln
        self.cell(0, 10, f"Appeal for {self.doc_id}", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
        self.set_font("Helvetica", "", 9)
        self.cell(0, 10, f"Generated {self.generated_at}", align="C")
        self.ln(5)

    def footer(self) -> None:  # type: ignore[override]
        self.set_y(-15)
        self.set_font("Helvetica", "", 9)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")


def build_appeal_pdf(
    doc_id: str,
    *,
    tone: str = "polite",
    audience: str = "payer",
    psl_delta: float | None = None,
    session_id: str | None = None,
) -> bytes:
    """Generate a PDF representation of the appeal letter."""

    if not _HAS_FPDF:  # pragma: no cover - requires fpdf2
        raise RuntimeError(
            "fpdf (fpdf2) is required to build a PDF appeal. Install fpdf2 and try again."
        )

    context = _build_context(doc_id, tone, audience, session_id)
    breakdown_summary = context["breakdown_summary"]
    risk_summary = context["risk_summary"]
    selected_tone = context["selected_tone"]
    greeting = context["greeting"]

    # use timezone-aware UTC timestamp and normalize to trailing Z
    timestamp = datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")
    pdf = _AppealPDF(doc_id=context["doc_id"], generated_at=timestamp)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, context["subject"], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, f"Tone: {tone.title()} | Audience: {audience.title()}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    _safe_multi = None
    try:
        # define a small wrapper that falls back gracefully if multi_cell fails
        def _safe_multi_cell(pdf_obj, w, h, txt):
            try:
                pdf_obj.multi_cell(w, h, txt)
            except FPDFException:
                # try shrinking font a couple of times before giving up
                orig_size = pdf_obj.font_size_pt
                for _ in range(3):
                    try:
                        pdf_obj.set_font(pdf_obj.font_family, pdf_obj.font_style, max(6, int(orig_size) - 2))
                        pdf_obj.multi_cell(w, h, txt)
                        return
                    except FPDFException:
                        continue
                # final fallback: write a truncated single-line cell
                safe = txt.replace("\n", " ")[:200]
                pdf_obj.cell(0, h, safe)

        _safe_multi = _safe_multi_cell
    except Exception:
        _safe_multi = lambda pdf_obj, w, h, txt: pdf_obj.multi_cell(w, h, txt)

    _safe_multi(pdf, 0, 6, f"{greeting},")
    pdf.ln(2)
    _safe_multi(pdf, 0, 6, selected_tone["opening"])
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Primary math evidence (PME) summary", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 11)
    if breakdown_summary:
        for line in breakdown_summary:
            _safe_multi(pdf, 0, 6, f"- {line}")
    else:
        pdf.multi_cell(0, 6, "No structured totals available.")
    pdf.ln(2)

    if risk_summary:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Risk considerations", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 11)
        for line in risk_summary:
            _safe_multi(pdf, 0, 6, f"- {line}")
        pdf.ln(2)

    if psl_delta is not None:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Policy simulation variance", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 11)
        _safe_multi(
            pdf,
            0,
            6,
            f"Expected vs billed patient responsibility delta: {_fmt_currency(psl_delta)}",
        )
        pdf.ln(2)

    pdf.set_font("Helvetica", "", 11)
    _safe_multi(pdf, 0, 6, selected_tone["closing"])
    pdf.ln(4)
    _safe_multi(pdf, 0, 6, "Sincerely,")
    _safe_multi(pdf, 0, 6, "Patient Advocate")
    pdf.ln(4)

    exhibit_index = _build_exhibit_index(bool(risk_summary))
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Exhibits", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 11)
    for exhibit in exhibit_index:
        label = exhibit["label"]
        title = exhibit["title"]
        description = exhibit.get("description", "")
        _safe_multi(pdf, 0, 6, f"{label}. {title} - {description}")

    pdf_bytes = pdf.output(dest="S")
    return bytes(pdf_bytes)
