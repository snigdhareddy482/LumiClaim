"""PDF/DOCX/Image extraction helpers.

Provides parse_pdf, parse_docx, parse_image and a small normalizer that
maps common table headers to the project schema. Optional dependencies
are used when available and the functions degrade gracefully.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import List, Tuple, Dict, Any

from backend.llm_adapters.gemini_adapter import extract_text_from_image

try:
    import pdfplumber  # type: ignore
except Exception:  # pragma: no cover - optional
    pdfplumber = None

try:
    import camelot  # type: ignore
except Exception:  # pragma: no cover - optional
    camelot = None

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover - optional
    Document = None

try:
    from PIL import Image  # type: ignore
    import pytesseract  # type: ignore
except Exception:  # pragma: no cover - optional
    Image = None
    pytesseract = None


SchemaRow = Dict[str, Any]


def _map_header(h: str) -> str:
    h = (h or "").strip().lower()
    mapping = {
        "cpt": "cpt",
        "procedure code": "cpt",
        "code": "cpt",
        "procedure": "cpt",
        "modifier": "modifier",
        "mod": "modifier",
        "billed": "billed",
        "amount billed": "billed",
        "charge": "billed",
        "allowed": "allowed",
        "allowed amount": "allowed",
        "insurer paid": "insurer_paid",
        "paid": "insurer_paid",
        "patient responsibility": "patient_resp",
        "patient resp": "patient_resp",
        "adjustments": "adjustments",
        "adj": "adjustments",
        "description": "description",
    }
    for key, val in mapping.items():
        if key in h:
            return val
    return h.replace(" ", "_")


def _fabricate_cell_id(table_idx: int, row_idx: int, col_idx: int) -> str:
    return f"tbl{table_idx}:R{row_idx}C{col_idx}"


def _normalize_table(table_rows: List[List[str]], page: int, table_idx: int) -> List[SchemaRow]:
    """Normalize a table (list-of-rows) into the schema rows.

    Expects first row to be headers. If headers are missing, uses positional
    guesses.
    """
    rows: List[SchemaRow] = []
    if not table_rows:
        return rows

    headers = table_rows[0]
    mapped = [(_map_header(h), idx) for idx, h in enumerate(headers)]

    for r_idx, raw_row in enumerate(table_rows[1:], start=1):
        out: SchemaRow = {
            "line_id": None,
            "page": page,
            "cell_id": _fabricate_cell_id(table_idx, r_idx, 0),
            "cpt": None,
            "modifier": None,
            "billed": None,
            "allowed": None,
            "insurer_paid": None,
            "adjustments": [],
            "patient_resp": None,
            "description": None,
        }
        for col_idx, cell in enumerate(raw_row):
            if col_idx >= len(mapped):
                # if there are more columns than headers, fabricate a key
                key = f"col_{col_idx}"
            else:
                key = mapped[col_idx][0]

            val = cell.strip() if isinstance(cell, str) else cell
            if key == "adjustments":
                # split on common separators
                parts = [p.strip() for p in str(val).split(";") if p.strip()]
                out["adjustments"] = parts
            elif key in {"billed", "allowed", "insurer_paid", "patient_resp"}:
                # try to parse numeric remove $ and commas
                try:
                    num = float(str(val).replace("$", "").replace(",", ""))
                except Exception:
                    num = None
                out[key] = num
            else:
                out[key] = val

        # ensure cell_id contains row index
        out["cell_id"] = out.get("cell_id") or _fabricate_cell_id(table_idx, r_idx, 0)
        out["line_id"] = f"L-{page}-{table_idx}-{r_idx}"
        rows.append(out)

    return rows


def parse_pdf(path: str) -> Tuple[List[SchemaRow], List[str], List[str]]:
    """Extract per-page raw text and attempt to parse tables into schema rows.

    Returns (rows, raw_text_pages, notes).
    """
    p = Path(path)
    rows: List[SchemaRow] = []
    raw_pages: List[str] = []
    notes: List[str] = []

    if pdfplumber is None:
        notes.append("pdfplumber unavailable; no text extracted")
    else:
        try:
            with pdfplumber.open(p) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text() or ""
                    
                    # Fallback to Gemini OCR if text is sparse (likely scanned)
                    if len(text.strip()) < 50 and hasattr(page, "to_image"):
                        try:
                            # Resolution 300 is standard for OCR
                            img = page.to_image(resolution=300).original
                            ocr_text = extract_text_from_image(img)
                            if ocr_text and len(ocr_text) > len(text):
                                text = ocr_text
                                notes.append(f"page_{i}_ocr_gemini")
                        except Exception as e:
                            notes.append(f"ocr_failed_page_{i}:{e}")
                            
                    raw_pages.append(text)
        except Exception as exc:
            notes.append(f"pdfplumber_failed:{exc}")

    # Try tables with camelot if present
    if camelot is None:
        notes.append("camelot unavailable; skipping table extraction")
    else:
        try:
            tables = camelot.read_pdf(str(p), pages="all")
            for t_idx, table in enumerate(tables):
                # table.df is pandas DataFrame; convert to list-of-lists of strings
                data = table.df.fillna("").astype(str).values.tolist()
                page_no = int(table.parsing_report.get("page", 1)) if hasattr(table, "parsing_report") else 1
                norm = _normalize_table(data, page_no, t_idx)
                rows.extend(norm)
        except Exception as exc:
            notes.append(f"camelot_failed:{exc}")

    # Fallback: if no structured rows found, try heuristic on the raw text
    if not rows and raw_pages:
        for p_idx, p_text in enumerate(raw_pages, start=1):
            heur_rows = _extract_rows_from_text(p_text, page=p_idx)
            if heur_rows:
                rows.extend(heur_rows)
        if rows:
            notes.append(f"heuristic_pdf: found {len(rows)} rows from text")

    return rows, raw_pages, notes


def parse_docx(path: str) -> Tuple[List[SchemaRow], List[str], List[str]]:
    """Extract text and tables from a DOCX file.

    Returns (rows, raw_text_pages, notes).
    """
    p = Path(path)
    rows: List[SchemaRow] = []
    raw_pages: List[str] = []
    notes: List[str] = []

    if Document is None:
        notes.append("python-docx unavailable; cannot parse docx")
        return rows, raw_pages, notes

    try:
        doc = Document(str(p))
    except Exception as exc:
        notes.append(f"docx_open_failed:{exc}")
        return rows, raw_pages, notes

    # collect paragraphs as a single 'page' (docx has no pages)
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    raw_pages.append("\n".join(paras))

    # process tables
    for t_idx, table in enumerate(doc.tables):
        data: List[List[str]] = []
        for r in table.rows:
            data.append([c.text for c in r.cells])
        norm = _normalize_table(data, 1, t_idx)
        rows.extend(norm)

    return rows, raw_pages, notes


    return rows, raw_pages, notes


def parse_image(path: str) -> Tuple[List[SchemaRow], List[str], List[str]]:
    """OCR an image into text pages and return empty rows.

    Returns (rows, raw_text_pages, notes).
    """
    rows: List[SchemaRow] = []
    raw_pages: List[str] = []
    notes: List[str] = []

    try:
        from PIL import Image as PILImage
        img = PILImage.open(path)
    except Exception as e:
        notes.append(f"image_open_failed: {e}")
        return rows, raw_pages, notes

    text = ""
    
    # Try Gemini OCR first as it is stronger than Tesseract for complex docs
    try:
        text = extract_text_from_image(img)
        if text:
            notes.append("gemini_ocr_success")
    except Exception as e:
        notes.append(f"gemini_ocr_failed: {e}")
    
    # Fallback to Tesseract if Gemini fails or is not configured
    if not text and pytesseract is not None:
        try:
            text = pytesseract.image_to_string(img)
            if text:
                notes.append("tesseract_ocr_success")
        except Exception as e:
            notes.append(f"tesseract_failed: {e}")
            
    raw_pages.append(text or "")
    
    # heuristic: try to extract rows from the OCR text
    if text:
        extracted_rows = _extract_rows_from_text(text, page=1)
        if extracted_rows:
            rows.extend(extracted_rows)
            notes.append(f"heuristic_ocr: found {len(extracted_rows)} rows")
        else:
            notes.append("heuristic_ocr: no structured rows found")

    return rows, raw_pages, notes


def _extract_rows_from_text(text: str, page: int = 1) -> List[SchemaRow]:
    """Heuristic extraction of EOB rows from raw text lines.
    
    Looks for lines starting with dates or containing multiple dollar amounts.
    """
    import re
    rows: List[SchemaRow] = []
    
    lines = text.splitlines()
    # Pattern: Date ... CPT? ... Money ... Money
    # 02/24/02/24  chiropractmanj 1-2regions  $40.00  $0.00  13106  $3.77  $36.23
    # We'll be lenient: look for at least 2 distinct dollar amounts or numbers
    
    # Matches MM/DD/YY or MM/DD/YYYY
    date_pattern = re.compile(r'(\d{1,2}/\d{1,2}/\d{2,4})')
    # Matches money $40.00 or 40.00
    money_pattern = re.compile(r'(\$?\d{1,3}(?:,\d{3})*\.\d{2})')
    
    for idx, line in enumerate(lines):
        if not line.strip():
            continue
            
        # 1. Check for date
        date_match = date_pattern.search(line)
        
        # 2. Check for monies
        monies = money_pattern.findall(line)
        
        # Candidates for valid claim lines usually have a date AND at least 2 money figures (Billed, Allowed, Paid, etc)
        # Or just many money figures
        if len(monies) >= 2:
            row: SchemaRow = {
                "line_id": f"L-{page}-ocr-{idx}",
                "page": page,
                "cell_id": f"ocr:R{idx}",
                "cpt": None,
                "modifier": None,
                "billed": None,
                "allowed": None,
                "insurer_paid": None,
                "adjustments": [],
                "patient_resp": None,
                "description": line.strip(), # fallback description is whole line
            }
            
            # Parsing logic:
            # We assume monies are in order: Billed, [Ineligible], [Discount], [Deductible], [Copay], [Coins], Paid, [Resp]
            # This is highly variable. We'll try to map largest to Billed.
            # And usually the last non-zero might be Patient Resp? 
            # This is tricky without spatial layout.
            
            # Simple heuristic: 
            # Max value -> Billed
            # If we detect "Patient Responsibility" in headers we might know index, but here we don't.
            
            # Let's clean values to floats
            floats = []
            for m in monies:
                try:
                    f = float(m.replace("$", "").replace(",", ""))
                    floats.append(f)
                except:
                    pass
            
            if not floats:
                continue

            # Assign Billed (usually the first or max)
            # In EOBs: Date, Proc, Billed, Allowed, ...
            # The first money is often Billed.
            row["billed"] = floats[0]
            
            # If we have >= 3 values, maybe Billed, Allowed, Paid?
            # Or Billed, Ineligible, Allowed?
            # Let's try to find Patient Responsibility. It's often the last column.
            if len(floats) >= 3:
                row["patient_resp"] = floats[-1]
                row["insurer_paid"] = floats[-2] if len(floats) > 3 else 0.0
                row["allowed"] = floats[1] # Guessing
            elif len(floats) == 2:
                # Billed, Paid? or Billed, Resp?
                row["allowed"] = floats[0]
                row["patient_resp"] = floats[1]

            # Try to grab CPT? 5 digit code.
            # 13106
            cpt_match = re.search(r'\b\d{5}\b', line)
            if cpt_match:
                row["cpt"] = cpt_match.group(0)
            
            rows.append(row)
            
    return rows
