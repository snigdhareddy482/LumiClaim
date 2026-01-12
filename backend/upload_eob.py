"""Handlers and helpers for multipart EOB uploads.

This module tries to extract text from uploaded files using pure-Python
libraries if available and degrades gracefully when optional packages
or external OCR binaries are missing.
"""

from __future__ import annotations

import json
import os
import secrets
import shutil
import string
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

try:
    from docx import Document  # python-docx
except Exception:  # pragma: no cover - optional
    Document = None

try:  # PDF extraction if available
    import PyPDF2  # type: ignore
except Exception:  # pragma: no cover - optional
    PyPDF2 = None

try:  # image OCR support (optional)
    from PIL import Image  # type: ignore
    import pytesseract  # type: ignore
except Exception:  # pragma: no cover - optional
    Image = None
    pytesseract = None

from fastapi import HTTPException

from backend import extractors
from backend.session import (
    SESSION_ROOT,
    append_claim_rows,
    append_raw_pages,
    ensure_session_dirs,
    ensure_session_files,
)
from backend.llm_adapters.groq_adapter import extract_structured_eob_text, QuotaExceededError, extract_text_from_image


BASE_SESSION_PATH = SESSION_ROOT

MAX_BYTES = 15 * 1024 * 1024  # 15 MB

ALLOWED_EXT = {".pdf", ".docx", ".png", ".jpg", ".jpeg"}


def _make_session_dirs(session_id: str) -> Tuple[Path, Path]:
    ensure_session_files(session_id)
    _, raw, extracted = ensure_session_dirs(session_id)
    return raw, extracted


def _safe_doc_id() -> str:
    """Return a sequential persistent doc id in the form EOB-###.

    Stores a tiny JSON file under data/user_sessions/_counter.json with
    structure {counter, created_by, created_at, updated_at}. If an older
    _counter.txt exists (legacy), attempt to migrate it.
    """
    import datetime

    counter_file = BASE_SESSION_PATH / "_counter.json"
    # migrate legacy text file
    legacy = BASE_SESSION_PATH / "_counter.txt"
    try:
        if legacy.exists() and not counter_file.exists():
            try:
                raw = legacy.read_text(encoding="utf-8").strip()
                n = int(raw)
            except Exception:
                n = 0
            payload = {
                "counter": max(1, n),
                "created_by": "migrated",
                # timezone-aware timestamps
                "created_at": datetime.datetime.now(datetime.timezone.utc).isoformat().replace("+00:00", "Z"),
                "updated_at": datetime.datetime.now(datetime.timezone.utc).isoformat().replace("+00:00", "Z"),
            }
            counter_file.write_text(json.dumps(payload), encoding="utf-8")

        if not counter_file.exists():
            payload = {
                "counter": 1,
                "created_by": "system",
                "created_at": datetime.datetime.now(datetime.timezone.utc).isoformat().replace("+00:00", "Z"),
                "updated_at": datetime.datetime.now(datetime.timezone.utc).isoformat().replace("+00:00", "Z"),
            }
            counter_file.write_text(json.dumps(payload), encoding="utf-8")
            num = 1
        else:
            raw = counter_file.read_text(encoding="utf-8").strip()
            data = {}
            try:
                data = json.loads(raw)
                num = int(data.get("counter", 0)) + 1
            except Exception:
                num = secrets.randbelow(1000)

            created_by = "system"
            created_at = datetime.datetime.now(datetime.timezone.utc).isoformat().replace("+00:00", "Z")
            try:
                if isinstance(data, dict):
                    created_by = data.get("created_by", created_by)
                    created_at = data.get("created_at", created_at)
            except Exception:
                pass

            payload = {
                "counter": int(num),
                "created_by": created_by,
                "created_at": created_at,
                "updated_at": datetime.datetime.now(datetime.timezone.utc).isoformat().replace("+00:00", "Z"),
            }
            counter_file.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    except Exception:
        # fallback to random when filesystem is unavailable
        num = secrets.randbelow(1000)

    return f"EOB-{int(num):03d}"


def _preview_snippet(text: str, length: int = 200) -> str:
    s = " ".join(text.strip().split())
    return s[:length]


def _detect_file_ext(filename: str) -> str:
    return Path(filename).suffix.lower()


def _save_raw_file(raw_dir: Path, filename: str, content: bytes) -> Path:
    target = raw_dir / filename
    with open(target, "wb") as fh:
        fh.write(content)
    return target


def _save_extracted_json(extracted_dir: Path, doc_id: str, payload: Dict[str, Any]) -> None:
    path = extracted_dir / f"{doc_id}.json"
    with open(path, "w", encoding="utf-8") as fh:
        json.dump(payload, fh, ensure_ascii=False, indent=2)


def _extract_from_docx_bytes(content: bytes) -> Tuple[str, int]:
    if Document is None:
        raise RuntimeError("python-docx not installed")
    from io import BytesIO

    buf = BytesIO(content)
    doc = Document(buf)
    text_parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n".join(text_parts)
    # docx doesn't have pages in this context; report paragraphs as pages approximation
    pages = max(1, len(doc.paragraphs) // 40)
    return text, pages


def _extract_from_pdf_bytes(content: bytes) -> Tuple[str, int]:
    if PyPDF2 is None:
        raise RuntimeError("PyPDF2 not installed")
    from io import BytesIO

    reader = PyPDF2.PdfReader(BytesIO(content))
    texts = []
    for page in reader.pages:
        try:
            texts.append(page.extract_text() or "")
        except Exception:
            texts.append("")
    pages = len(reader.pages)
    combined_text = "\n".join(texts)
    
    # If PyPDF2 got very little text, the PDF is likely scanned images.
    # Try Gemini Vision OCR on the first page as fallback.
    if len(combined_text.strip()) < 50:
        try:
            from pdf2image import convert_from_bytes
            from backend.llm_adapters.gemini_adapter import extract_text_from_image
            
            # Convert first page to image
            images = convert_from_bytes(content, first_page=1, last_page=1)
            if images:
                ocr_text = extract_text_from_image(images[0])
                if ocr_text and len(ocr_text.strip()) > 20:
                    print("[PDF] Gemini OCR fallback succeeded")
                    return ocr_text, pages
        except QuotaExceededError:
            raise  # bubble up to handler
        except Exception as e:
            print(f"[PDF] Gemini OCR fallback failed: {e}")
    
    return combined_text, pages



def _extract_from_image_bytes(content: bytes) -> Tuple[str, int]:
    from io import BytesIO
    from PIL import Image as PILImage
    
    img = PILImage.open(BytesIO(content))
    
    # Try Gemini Vision OCR first (most powerful)
    try:
        from backend.llm_adapters.gemini_adapter import extract_text_from_image
        text = extract_text_from_image(img)
        if text and len(text.strip()) > 20:
            return text, 1
    except QuotaExceededError:
        raise
    except Exception as e:
        # Log but don't fail - try Tesseract next
        print(f"Gemini OCR failed: {e}")
    
    # Fallback to Tesseract if available
    if pytesseract is not None:
        try:
            text = pytesseract.image_to_string(img)
            if text and len(text.strip()) > 20:
                return text, 1
        except Exception as e:
            print(f"Tesseract OCR failed: {e}")
    
    # Return empty if both failed
    return "", 1


def handle_upload_file(filename: str, content: bytes, session_id: str | None = None) -> Dict[str, Any]:
    """Validate, save the raw file, attempt extraction, redact, and save artifacts.

    Returns a dict with session_id, doc_id, file_type, pages, notes, preview.
    Raises HTTPException for friendly errors.
    """

    if not filename:
        raise HTTPException(status_code=400, detail="filename is required")

    ext = _detect_file_ext(filename)
    if ext not in ALLOWED_EXT:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")

    if len(content) == 0:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")
    if len(content) > MAX_BYTES:
        raise HTTPException(status_code=400, detail="File too large; maximum allowed is 15MB")

    if session_id is None:
        raise HTTPException(status_code=400, detail="No session. Call /session/start first.")

    raw_dir, extracted_dir = _make_session_dirs(session_id)
    
    # -- Deduplication Start --
    import hashlib
    file_hash = hashlib.sha256(content).hexdigest()
    hashes_file = BASE_SESSION_PATH / session_id / "_hashes.json"
    hashes_map = {}
    if hashes_file.exists():
        try:
            hashes_map = json.loads(hashes_file.read_text(encoding="utf-8"))
        except Exception:
            pass

    if file_hash in hashes_map:
        # Return existing document metadata
        existing_doc_id = hashes_map[file_hash]
        # Try to load existing artifact to return proper preview
        try:
            existing_path = extracted_dir / f"{existing_doc_id}.json"
            if existing_path.exists():
                artifact = json.loads(existing_path.read_text(encoding="utf-8"))
                artifact["notes"] = artifact.get("notes", []) + ["Duplicate upload detected; returned existing document."]
                return {
                    "session_id": session_id,
                    "doc_id": existing_doc_id,
                    "preview": {
                        "rows": [], # We don't reload rows here to save time, or we could load from claims_struct
                        "text_snippets": [artifact.get("extracted_text_preview", "")],
                    },
                    "duplicate": True
                }
        except Exception:
            # If artifact missing, proceed to re-process but maybe warn?
            # Ideally re-process if missing. For now, let's fall through to re-process if artifact header missing.
            pass
    # -- Deduplication End --

    safe_name = "".join(ch for ch in filename if ch in (string.ascii_letters + string.digits + ".-_"))
    saved_path = _save_raw_file(raw_dir, safe_name or filename, content)

    notes_parts = []
    extracted_text = ""
    pages = 0
    parsed_rows: list = []
    raw_pages: list = []

    # Try format-specific extraction (text layer) and parsers for structured rows
    try:
        if ext == ".docx":
            if Document is None:
                notes_parts.append("DOCX parsing unavailable")
            else:
                extracted_text, pages = _extract_from_docx_bytes(content)
                # parse docx tables/text for structured rows
                try:
                    parsed_rows, raw_pages, parse_notes = extractors.parse_docx(str(saved_path))
                    notes_parts.extend(parse_notes or [])
                except Exception as _:
                    parsed_rows, raw_pages = [], []
        elif ext == ".pdf":
            if PyPDF2 is None:
                notes_parts.append("PDF text extraction unavailable")
            else:
                extracted_text, pages = _extract_from_pdf_bytes(content)
                try:
                    parsed_rows, raw_pages, parse_notes = extractors.parse_pdf(str(saved_path))
                    notes_parts.extend(parse_notes or [])
                except Exception as e:
                    try:
                        with open("backend_crash.log", "a") as f:
                            f.write(f"CRASH: {e}\n")
                            import traceback
                            traceback.print_exc(file=f)
                    except:
                        pass
                    parsed_rows, raw_pages = [], []
        elif ext in {".png", ".jpg", ".jpeg"}:
            # try OCR
            extracted_text, pages = _extract_from_image_bytes(content)
            if extracted_text:
                notes_parts.append("ocr_success")
            try:
                parsed_rows, raw_pages, parse_notes = extractors.parse_image(str(saved_path))
                notes_parts.extend(parse_notes or [])
            except Exception:
                parsed_rows, raw_pages = [], []
    except QuotaExceededError:
        print("[Upload] Gemini Quota Exceeded (OCR/Text Layer)")
        notes_parts.append("quota_exceeded_try_later")
    except Exception as exc:  # pragma: no cover - defensive
        notes_parts.append(f"extraction_failed: {str(exc)}")

    
    # Consolidate text: if extracted_text is missing or too short, and we have raw pages
    is_text_usable = extracted_text and len(extracted_text.strip()) > 20
    if not is_text_usable and raw_pages:
        extracted_text = "\n".join(str(p) for p in raw_pages)

    # Fallback: if we have text but no structured rows, try Gemini Text Extraction
    if not parsed_rows and extracted_text and len(extracted_text.strip()) > 20:
        try:
            ai_rows = extract_structured_eob_text(extracted_text)
            if ai_rows:
                parsed_rows = []
                for item in ai_rows:
                    parsed_rows.append({
                        "line_id": f"L-gemini-fallback-{len(parsed_rows)}",
                        "page": 1,
                        "cell_id": "gemini:text_fallback",
                        "cpt": item.get("cpt"),
                        "modifier": None,
                        "billed": item.get("billed"),
                        "allowed": item.get("allowed"),
                        "insurer_paid": item.get("insurer_paid"),
                        "adjustments": item.get("adjustments", []),
                        "patient_resp": item.get("patient_resp"),
                        "description": item.get("description"),
                    })
                notes_parts.append(f"gemini_text_fallback_success: {len(parsed_rows)} rows")
            else:
                notes_parts.append(f"gemini_text_fallback_empty_result_len_{len(extracted_text)}")
        except QuotaExceededError:
            print("[Upload] Gemini Quota Exceeded during fallback.")
            notes_parts.append("quota_exceeded_try_later")
        except Exception as e:
            # log but don't fail upload
            print(f"Gemini fallback failed: {e}")
            notes_parts.append(f"gemini_text_fallback_failed: {str(e)}")

    # If we failed to extract any text, keep minimal placeholder
    if not extracted_text:
        extracted_text = ""

    # Redact PHI
    redacted = redact_text(extracted_text)

    # Save extracted artifact
    doc_id = _safe_doc_id()
    artifact: Dict[str, Any] = {
        "doc_id": doc_id,
        "session_id": session_id,
        "filename": safe_name,
        "file_type": ext.lstrip("."),
        "pages": pages,
        "notes": notes_parts,
        "extracted_text_preview": _preview_snippet(redacted),
    }

    # Persist extracted text and metadata
    _save_extracted_json(extracted_dir, doc_id, {**artifact, "redacted_text": redacted})

    # If parser produced structured rows/raw pages, append them to per-session files
    sanitized_rows: list[dict[str, Any]] = []
    sanitized_pages: list[dict[str, Any]] = []

    try:
        claims_struct_path = extracted_dir / "claims_struct.json"
        claims_raw_path = extracted_dir / "claims_raw.json"

        parsed_rows = locals().get("parsed_rows", []) or []
        raw_pages = locals().get("raw_pages", []) or []

        for row in parsed_rows:
            if not isinstance(row, dict):
                continue
            cleaned = dict(row)
            cleaned["doc_id"] = doc_id
            for key, value in list(cleaned.items()):
                if isinstance(value, str) and value.strip():
                    cleaned[key] = redact_text(value)
            sanitized_rows.append(cleaned)

        existing_struct = []
        if claims_struct_path.exists():
            try:
                existing_struct = json.loads(claims_struct_path.read_text(encoding="utf-8"))
                if not isinstance(existing_struct, list):
                    existing_struct = []
            except Exception:
                existing_struct = []

        if sanitized_rows:
            existing_struct.extend(sanitized_rows)
            claims_struct_path.write_text(json.dumps(existing_struct, ensure_ascii=False, indent=2), encoding="utf-8")
            append_claim_rows(session_id, sanitized_rows)

        existing_raw = []
        if claims_raw_path.exists():
            try:
                existing_raw = json.loads(claims_raw_path.read_text(encoding="utf-8"))
                if not isinstance(existing_raw, list):
                    existing_raw = []
            except Exception:
                existing_raw = []

        for index, txt in enumerate(raw_pages, start=1):
            snippet = redact_text(str(txt))
            entry = {"doc_id": doc_id, "page": index, "text": snippet}
            existing_raw.append(entry)
            sanitized_pages.append(entry)

        if sanitized_pages:
            claims_raw_path.write_text(json.dumps(existing_raw, ensure_ascii=False, indent=2), encoding="utf-8")
            append_raw_pages(session_id, sanitized_pages)
    except Exception:
        # non-fatal; do not block the upload if writing artifacts fails
        pass

    rows_preview = sanitized_rows[:3]

    text_snippets: list[str] = []
    for entry in sanitized_pages[:2]:
        snippet = _preview_snippet(str(entry.get("text", "")), length=200)
        if snippet:
            text_snippets.append(snippet)

    if not text_snippets:
        text_snippets.append(artifact["extracted_text_preview"])

    # -- Deduplication Persist --
    try:
        hashes_map[file_hash] = doc_id
        hashes_file.write_text(json.dumps(hashes_map, indent=2), encoding="utf-8")
    except Exception:
        pass
    # -- End Persist --

    return {
        "session_id": session_id,
        "doc_id": doc_id,
        "filename": safe_name,
        "notes": notes_parts,
        "extracted_text_preview": artifact.get("extracted_text_preview"),
        "preview": {
            "rows": rows_preview,
            "text_snippets": text_snippets,
        },
    }
