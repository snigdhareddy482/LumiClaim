from __future__ import annotations

from pathlib import Path

import pytest

from backend import extractors


def test_parse_docx(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    from docx import Document

    doc = Document()
    doc.add_paragraph("Header paragraph for testing.")
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "CPT"
    hdr_cells[1].text = "Billed"
    hdr_cells[2].text = "Allowed"
    hdr_cells[3].text = "Patient Responsibility"

    # add two rows
    r = table.add_row().cells
    r[0].text = "12345"
    r[1].text = "$100.00"
    r[2].text = "$80.00"
    r[3].text = "$20.00"

    r2 = table.add_row().cells
    r2[0].text = "67890"
    r2[1].text = "$200.00"
    r2[2].text = "$160.00"
    r2[3].text = "$40.00"

    path = tmp_path / "sample.docx"
    doc.save(str(path))

    rows, raw_pages, notes = extractors.parse_docx(str(path))

    assert isinstance(rows, list)
    # two data rows expected
    assert len(rows) == 2
    assert any("12345" in str(r.get("cpt") or r.get("col_0") or "") for r in rows)
    assert raw_pages and "Header paragraph" in raw_pages[0]


def test_parse_pdf(tmp_path: Path) -> None:
    # pdfplumber is optional; skip if not present
    pytest.importorskip("pdfplumber")
    from fpdf import FPDF, XPos, YPos

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    # use new positioning args to avoid deprecated `ln` parameter
    pdf.cell(0, 10, "Amount Billed: $123.45", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.cell(0, 10, "Allowed Amount: $100.00", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    path = tmp_path / "sample.pdf"
    pdf.output(str(path))

    rows, raw_pages, notes = extractors.parse_pdf(str(path))

    # raw_pages should capture the text layer when pdfplumber is available
    assert isinstance(raw_pages, list)
    assert any("Amount Billed" in (p or "") for p in raw_pages)


def test_parse_image(tmp_path: Path) -> None:
    pytest.importorskip("pytesseract")
    pytest.importorskip("PIL")
    import pytesseract
    from PIL import Image, ImageDraw, ImageFont

    # skip if tesseract binary not available
    try:
        _ = pytesseract.get_tesseract_version()
    except Exception:
        pytest.skip("tesseract binary not available; skipping OCR test")

    img = Image.new("RGB", (200, 80), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    # default font
    d.text((10, 10), "Hello 123", fill=(0, 0, 0))

    path = tmp_path / "sample.png"
    img.save(path)

    rows, raw_pages, notes = extractors.parse_image(str(path))

    # OCR should produce some text (if tesseract is present)
    assert isinstance(raw_pages, list)
    assert raw_pages and any("Hello" in (p or "") for p in raw_pages)
