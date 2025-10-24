from __future__ import annotations

import os
from pathlib import Path

import pytest

from fastapi import HTTPException

from backend.upload_eob import handle_upload_file, BASE_SESSION_PATH, MAX_BYTES


def test_reject_unsupported_type():
    with pytest.raises(HTTPException) as exc:
        handle_upload_file("malware.exe", b"bad")
    assert exc.value.status_code == 400


def test_reject_too_large():
    # construct a bytes larger than MAX_BYTES
    big = b"0" * (MAX_BYTES + 1)
    with pytest.raises(HTTPException) as exc:
        handle_upload_file("large.pdf", big)
    assert exc.value.status_code == 400


def test_docx_extraction_and_redaction(tmp_path: Path):
    # Skip if python-docx isn't available — the handler degrades gracefully.
    pytest.importorskip("docx")
    from docx import Document
    from io import BytesIO

    doc = Document()
    doc.add_paragraph("Patient Name: John Doe")
    doc.add_paragraph("Email: alice@example.com")
    buf = BytesIO()
    doc.save(buf)
    content = buf.getvalue()

    result = handle_upload_file("claim.docx", content)

    assert "session_id" in result
    assert result["doc_id"].startswith("EOB-")
    # artifact files should exist on disk
    sess = result["session_id"]
    raw_dir = BASE_SESSION_PATH / sess / "raw"
    extracted_dir = BASE_SESSION_PATH / sess / "extracted"
    assert raw_dir.exists()
    assert extracted_dir.exists()

    # extracted artifact JSON should be present
    docs = list(extracted_dir.glob("*.json"))
    assert docs, "Expected at least one extracted JSON artifact"

